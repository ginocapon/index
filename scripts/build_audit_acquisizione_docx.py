#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera DOCX (e opz. PDF) audit strategico acquisizione incarichi."""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT_DOCX = ROOT / "documenti" / "Audit-Strategico-Acquisizione-Incarichi-Righetto-2026-08-29.docx"
OUT_PDF = ROOT / "documenti" / "Audit-Strategico-Acquisizione-Incarichi-Righetto-2026-08-29.pdf"


def build_docx() -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h1(text: str) -> None:
        doc.add_heading(text, level=1)

    def h2(text: str) -> None:
        doc.add_heading(text, level=2)

    def h3(text: str) -> None:
        doc.add_heading(text, level=3)

    def p(text: str, bold: bool = False) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold

    def bullet(items: list[str]) -> None:
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    def table(headers: list[str], rows: list[list[str]]) -> None:
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row):
                t.rows[ri + 1].cells[ci].text = cell
        doc.add_paragraph()

    # --- Cover ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("AUDIT STRATEGICO AVVERSARIALE\nDEL SITO")
    r.bold = True
    r.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Righetto Immobiliare — righettoimmobiliare.it")
    sr.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Verifica critica della capacità di acquisire nuovi immobili e nuovi proprietari\n")
    meta.add_run("Data: 29 agosto 2026\n")
    meta.add_run("Metodo: analisi codice, funnel, coda editoriale, GSC, pagine conversione")

    doc.add_page_break()

    h1("Risposta alla domanda centrale")
    p(
        "Il sito NON è ancora una macchina di acquisizione incarichi. Ha componenti utili ma "
        "scollegate (landing valutazione, servizio vendita/locazioni, ~15 pillar owner nel blog). "
        "Il traffico e la produzione editoriale recente sono orientati prevalentemente ad acquirenti "
        "e inquilini, non a proprietari che devono decidere se affidare un immobile. Non esiste "
        "misurazione del tasso di acquisizione incarichi: impossibile dimostrare che le modifiche "
        "recenti abbiano aumentato i mandati."
    )

    h1("A. Verdetto attuale — Dove siamo realmente")
    p(
        "Il sito è un ecosistema informativo maturo (132 articoli, FAQ, zone, servizi, form, GA4, "
        "lead tracking) con conversion path spezzato."
    )
    table(
        ["Cosa esiste", "Cosa manca"],
        [
            ["landing-valutazione.html — miglior pagina conversione", "Orfana: homepage non la linka"],
            ["servizio-vendita / servizio-locazioni — form dedicati", "Orfane dalla homepage e dal blog feed"],
            ["~15–16 pillar owner (mandato, costi vendere, home staging…)", "Stagnanti: ultimo pillar conversione 17 luglio 2026"],
            ["Form su ~91% articoli", "Form generico consulenza, non segmentato venditore/locatore"],
            ["Skill §16-TER/QUATER/QUINQUIES", "Ottimizza qualità contenuti, non funnel proprietari"],
            ["GSC: 271 click/28gg, brand forte", "Winner pages = affitto studenti, non venditori"],
        ],
    )

    h3("Percentuale sforzo editoriale collegata ad acquisizione incarichi (stimata)")
    table(
        ["Periodo", "A (diretto)", "B (indiretto)", "C+D (debole/dispersione)"],
        [
            ["Ultimi 30 articoli", "~20%", "~10%", "~70%"],
            ["Corpus intero (~132)", "~17–21%", "~11–14%", "~55–65%"],
            ["Batch 29 agosto (7 articoli)", "0% owner-conversion", "~15%", "~85%"],
        ],
    )
    p(
        "Il sistema non misura l'obiettivo. GA4 traccia lead_segmentation e speed_to_lead, ma non "
        "esiste un KPI «lead venditore/locatore → sopralluogo → mandato». Senza questo, ogni "
        "modifica SEO/editoriale resta non dimostrata sul piano commerciale."
    )

    h1("B. Errori principali")
    errors = [
        (
            "1. Funnel proprietario spezzato (errore strutturale)",
            "index.html — tutti i CTA owner puntano a vendere-casa-padova-errori, zero link a "
            "landing-valutazione, servizio-vendita, servizio-locazioni. Il percorso attuale aggiunge "
            "1–2 step inutili e usa una pagina con form più debole del landing dedicato.",
        ),
        (
            "2. Homepage dual-audience che favorisce l'acquirente",
            "Hero: «Vendere e comprare casa». Ricerca immobili, blocco mutuo, blog con ultimi "
            "articoli acquisto. Un proprietario trova segnali misti; la fiducia specialistica si diluisce.",
        ),
        (
            "3. Calendario editoriale invertito rispetto al business",
            "editorial-queue.json: 18 item published, 1 solo con conversion: true. Batch 29/08: "
            "7 articoli, 0 nuovi pillar owner. Si produce per NORMATIVA_FISCALE, ACQUISTO, MUTUI.",
        ),
        (
            "4. Claim numerici senza prova inline",
            "Incongruenze tempi vendita (60–90 gg vs 4,2 mesi); «+7%», «40% contatti», «10–15% da soli» "
            "senza fonte visibile; stima istantanea vs perizia 24h; mappa contatti placeholder.",
        ),
        (
            "5. CTA blog declassate nel template recente",
            "blog-costi-vendere-casa: 3× landing-valutazione (modello A). Articoli batch agosto 2026: "
            "solo consulenza generica. Solo 39% del corpus linka landing-valutazione.",
        ),
        (
            "6. GSC strategy cattura domanda sbagliata",
            "Batch Limena chiude gap comprare/affittare. blog-rendimento-affitto-padova: 139 impressioni, "
            "0 click. Nessuna keyword tracciata per vendere casa / valutazione immobile Padova.",
        ),
        (
            "7. Differenziazione non dimostrata",
            "Mancano casi vendita documentati, proof locale per comune, percorso naturale da contenuto "
            "owner a servizi distintivi (drone, tour 360).",
        ),
    ]
    for title, body in errors:
        h2(title)
        p(body)

    h1("C. Modifiche inutili o poco strategiche")
    table(
        ["MODIFICA", "OBIETTIVO", "CONTRIBUTO ACQUISIZIONE", "VERDETTO"],
        [
            ["§16-QUATER varietà visiva", "Evitare percezione articoli uguali", "Indiretto debole", "NON DIMOSTRATO"],
            ["§16-QUINQUIES FAQ/geo/proprietari", "FAQ + ciclo owner", "Potenziale, non eseguito", "MIGLIORARE"],
            ["§18 bonifica residui AI", "Pubblicabilità", "Higiene, non mandati", "MANTENERE"],
            ["§16-TER continuità sostanziale", "Anti-doppioni", "SEO, spinge acquisto/normativa", "RIVEDERE"],
            ["Batch condominio + proposta acquisto", "Test §16-TER", "Zero — acquirenti", "RIVEDERE"],
            ["Batch venerdì 29/08 (5 articoli)", "Comando editoriale", "~15% locatori", "RIVEDERE"],
            ["Batch Limena GSC", "Gap keyword locali", "Buyer/tenant traffic", "RIVEDERE"],
            ["4 articoli EN affitti/expat", "Visibilità internazionale", "Quasi zero mandati locali", "ELIMINARE"],
            ["lead-conversion.js A/B test", "Conversioni form", "Potenziale, funnel rotto", "MIGLIORARE"],
            ["Blog modal → contatti", "Intent da blog", "Destinazione generica", "MIGLIORARE"],
            ["Anti-doppioni SKIMM", "Qualità editoriale", "Operativo", "MANTENERE"],
            ["Foto AI, audit visivo", "Compliance + qualità", "Fiducia indiretta", "MANTENERE"],
        ],
    )

    h1("Percorso proprietario — 6 fasi")
    phases = [
        ("FASE 1 — Scoperta", "Parzialmente. Intercetta errori vendita, costi, mandato. NON intercetta head-term valutazione, vendere con mutuo, gestione delegata, decisione vendere vs affittare. ~70% produzione recente = traffico a bassa probabilità di mandato."),
        ("FASE 2 — Riconoscimento problema", "Funziona su vendere-casa-padova-errori e pillar owner. Non funziona su articoli macro (geopolitica, dazi) — zero ancoraggio al proprietario padovano."),
        ("FASE 3 — Utilità", "Utile: pillar con fonti GU/ADE/OMI. Generico/SEO: trend mercato, gergo, coliving — non spostano decisione di vendita."),
        ("FASE 4 — Fiducia", "Dimostrata: processo servizio-vendita, recensioni Google. Solo dichiarata: numeri performance. Danneggiata: incongruenze, stima vs perizia, mappa placeholder."),
        ("FASE 5 — Differenziazione", "Debole. Presenza locale e servizi tech esistono ma non sono percorso naturale da contenuto owner. Non distinguibile per cold owner da Google."),
        ("FASE 6 — Desiderio di contatto", "Form su servizi e landing valutazione. Si blocca: homepage → contenuto lungo; blog → consulenza generica. CTA non portano al form migliore."),
    ]
    for title, body in phases:
        h2(title)
        p(body)

    h1("Analisi pagine principali")
    table(
        ["Pagina", "Proprietario", "Problema compreso?", "Next step", "Motivo abbandono"],
        [
            ["Homepage", "Misto", "Parziale", "vendere-casa-errori", "Confusione audience, troppo buyer"],
            ["servizio-vendita", "Venditore", "Sì", "Form #richiedi", "Orfana da home"],
            ["servizio-locazioni", "Locatore", "Sì", "Form #richiedi", "Orfana"],
            ["landing-valutazione", "Pre-vendita", "Sì", "Form 3 step", "Orfana; stima vs perizia"],
            ["landing-consulenza", "Multi-intent", "Sì", "Form", "Vendita = 1 di 4 card"],
            ["vendere-casa-errori", "Venditore dubbioso", "Sì", "Email gate", "Non form migliore"],
            ["blog", "Vari", "Dipende", "Consulenza generica", "Feed recente = acquirenti"],
            ["faq", "Tutti", "Sì", "Contatti", "Skew mutuo/acquirente"],
            ["contatti", "Catch-all", "Generico", "Form", "No mappa reale"],
        ],
    )

    h1("Classificazione articoli (corpus ~132)")
    table(
        ["Classe", "% stimato", "Utilità acquisizione"],
        [
            ["A — Diretti (vendita/affitto/gestione)", "17–21%", "Alta se CTA verso valutazione"],
            ["B — Indiretti (trust, brand)", "11–14%", "Media"],
            ["C — Deboli (acquisto, mutuo)", "42–49%", "Bassa per mandati"],
            ["D — Dispersione (tenant, EN, macro)", "27–30%", "Quasi nulla"],
        ],
    )

    h1("Intenzioni di ricerca mancanti (owner)")
    table(
        ["Domanda proprietario", "Sito presente?", "Gap"],
        [
            ["Quanto vale casa Padova 2026", "Parziale", "No pillar head-term → landing-valutazione"],
            ["Come vendere casa Padova", "Parziale", "Refresh 2026 + CTA cluster"],
            ["Affittare casa (POV proprietario)", "Debole", "Pillar proprietario che affitta"],
            ["Cedolare secca / fiscalità locazione", "Parziale", "Pillar dedicato"],
            ["Vendere con mutuo in corso", "No", "Nuovo contenuto"],
            ["Vendere casa ereditata", "Parziale", "Non collegato a valutazione"],
            ["Gestione locazione delegata", "Servizio esiste", "Quasi zero contenuto"],
            ["Vendere/affittare Limena (owner)", "No", "Batch Limena = buyer"],
            ["Mandato non esclusivo vs esclusivo", "Solo esclusivo", "Confronto mancante"],
        ],
    )

    h1("Strategia AI — verifica critica")
    p(
        "Perché un'AI citerebbe Righetto vs ADE/Normattiva? Oggi: difficilmente su normativa pura; "
        "forse su guide locali Padova/Limena con OMI. Stesse fonti di chiunque — non moat. "
        "Costruire: report mercato trimestrale Padova/Limena con dati OMI + transazioni anonimizzate; "
        "FAQ owner ultra-specifiche per comune."
    )

    h1("Test di sincerità — 12 mesi con strategia attuale")
    p("Risposta: PARZIALMENTE", bold=True)
    p(
        "Il sito accumulerà autorevolezza e traffico informativo (acquirenti/inquilini), ma senza "
        "correzione funnel e calendario owner la probabilità di incremento significativo di nuovi "
        "incarichi resta limitata."
    )

    h1("Scenario negativo")
    h2("10 rischi concreti")
    bullet([
        "Homepage continua a mandare owner a contenuto educativo invece che a form",
        "Slot blog settimanale consumato da acquisto/normativa",
        "Claim numerici contestati → perdita fiducia",
        "GSC cresce su keyword buyer — successo illusorio",
        "landing-valutazione genera lead curiosi non venditori",
        "Competitor copia pillar owner con CTA più aggressive",
        "Saturazione contenuti simili — de-prioritization Google",
        "Tempo team su audit/skill vs funnel",
        "Chatbot/FAQ rispondono — utente non contatta",
        "Nessun CRM tag origine venditore",
    ])
    h2("10 errori strategici possibili")
    bullet([
        "Confondere articoli pubblicati con proprietari intercettati",
        "Investire in varietà visiva H2 prima di fix homepage CTA",
        "Batch Limena buyer quando serve owner Limena",
        "Refresh meta senza rewrite contenuto+CTA owner",
        "Pillar EN expat con risorse limitate",
        "Trend geopolitica come priorità venerdì",
        "Form generico consulenza su tutto il blog",
        "Ignorare servizio-gestione nel percorso owner affitto",
        "Non collegare zone-*.html a landing valutazione",
        "Misurare click GSC invece di lead venditore qualificati",
    ])
    h2("10 attività che fanno perdere tempo")
    bullet([
        "Audit visual variety su articoli buyer",
        "Generazione 8 WebP per articolo acquisto",
        "Bonifica §18 su file già ok",
        "Articoli gergo immobiliare (zero intent)",
        "Duplicati titoli in articoliStatici",
        "4 articoli EN rental per mercato locale",
        "Coliving/under-35 senza path owner",
        "Aggiornare memorie JSON senza cambiare coda",
        "SEO macro Italia quando priorità è Padova",
        "A/B test CTA su funnel non riparato",
    ])
    h2("10 ragioni per cui il traffico non diventa proprietari")
    bullet([
        "Intent query = acquisto, non vendita",
        "CTA blog → consulenza generica",
        "Contenuto risponde ma non propone passo successivo",
        "Proprietario legge 2500 parole senza CTA locale",
        "Form lungo vs WhatsApp one-tap competitor",
        "Manca urgenza mercato locale Q3 2026",
        "Testimonial non specifici per vendita",
        "Valutazione percepita come trucco (stima istantanea)",
        "Brand forte attira chi cerca Righetto, non cold owner",
        "Nessun retargeting owner",
    ])

    h1("Punteggi (0–100)")
    table(
        ["AREA", "PUNTEGGIO", "MOTIVAZIONE"],
        [
            ["Capacità di intercettare proprietari", "42", "Pillar esistono ma non alimentati; gap head-term"],
            ["Fiducia", "58", "Recensioni ok; claim incongruenti, mappa placeholder"],
            ["Autorevolezza", "72", "Volume, fonti, FAQ — alta per SEO, non convertita"],
            ["Competenza locale Padova/Veneto", "65", "Zone, OMI; manca proof transazioni"],
            ["Differenziazione", "38", "Servizi tech non posizionati come perché noi"],
            ["Trasformare traffico in contatti", "45", "Form ovunque ma routing sbagliato"],
            ["Capacità di acquisire incarichi", "40", "Infrastruttura sì, sistema integrato no"],
            ["Qualità contenuti", "78", "Lunghi, fontati, audit pass"],
            ["Utilità reale", "62", "Alta per info; media per decisione affidamento"],
            ["Strategia AI", "55", "AEO ok; manca contenuto distintivo"],
        ],
    )
    p("PUNTEGGIO COMPLESSIVO CAPACITÀ ACQUISIZIONE INCARICHI: 44/100", bold=True)

    h1("D. Opportunità mancanti")
    bullet([
        "Homepage owner-first — hero venditore/locatore, link servizi/valutazione",
        "Pillar «Valutazione casa Padova 2026» → landing-valutazione",
        "Coda editoriale 50/50 owner/buyer con conversion: true ogni 2 settimane",
        "Refresh blog-rendimento-affitto-padova (139 impr, 0 click)",
        "Owner batch Limena — vendere/affittare casa Limena (supply side)",
        "Case study vendite anonimizzate per comune",
        "Zone pages → CTA valutazione per quartiere",
        "KPI funnel GA4: lead venditore → mandato",
        "FAQ owner-first — ridurre skew mutuo",
        "servizio-gestione content hub per proprietari non residenti",
    ])

    h1("E. Priorità assolute (10 modifiche)")
    bullet([
        "Ricollegare homepage → landing-valutazione + servizio-vendita/locazioni",
        "Regola coda: 1 articolo owner conversion: true ogni ciclo",
        "Pillar valutazione + vendere casa Padova 2026 (pattern costi-vendere)",
        "Fix blog-rendimento-affitto-padova + servizio-locazioni",
        "Template blog Class A: obbligo landing-valutazione",
        "Modal blog homepage → landing-valutazione",
        "Allineare claim (tempi vendita, +7%) — fonte inline o rimuovere",
        "Contatti: mappa Google reale",
        "landing-valutazione: no opt-in pre-checked; chiarire stima vs perizia",
        "GSC dashboard owner — query venditore/locatore separate",
    ])

    h1("F. Elementi da non toccare")
    bullet([
        "servizio-vendita.html / servizio-locazioni.html — struttura e form",
        "landing-valutazione.html — UX form (fix etica, non stravolgere)",
        "landing-consulenza-immobiliare-gratuita.html",
        "Pillar owner esistenti (costi vendere, mandato, documenti vendita, home staging)",
        "Form lead Supabase + provenienza",
        "§18 publishability, anti-doppioni SKIMM",
        "Recensioni Google + schema AggregateRating",
        "Zone locali — collegare meglio, non rifare",
        "Tele/WhatsApp ripetuti",
    ])

    h1("G. Piano di correzione (ordine priorità)")
    steps = [
        "Fix routing CTA homepage → landing valutazione + servizi",
        "Aggiornare editorial-queue — prossimi 4 item: 2 owner + 2 refresh GSC",
        "Scrivere pillar valutazione 2026",
        "Refresh rendimento affitto",
        "Template blog Class A — CTA obbligatori",
        "Fix contatti mappa + claim incongruenti",
        "Owner batch Limena",
        "Case study su servizio-vendita",
        "GA4 dashboard owner funnel",
        "FAQ skew + link conversione",
        "Poi: varietà visiva §16-QUATER, articoli buyer, trend macro",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_page_break()
    h1("Verdetto finale obbligatorio")
    verdict = doc.add_paragraph()
    vr = verdict.add_run(
        "IL SITO È PARZIALMENTE ORIENTATO ALL'ACQUISIZIONE DI INCARICHI "
        "MA LA STRATEGIA ATTUALE PRESENTA LACUNE IMPORTANTI."
    )
    vr.bold = True
    vr.font.size = Pt(12)

    p(
        "Motivazione: esistono pagine conversione, pillar owner e form — quindi l'intenzione e parte "
        "dell'infrastruttura ci sono. Ma il sistema operativo recente (homepage, coda editoriale, "
        "batch GSC, ultimi 30 articoli) non collega traffico e contenuti all'acquisizione di "
        "proprietari. Le modifiche skill §16-TER/QUATER/QUINQUIES e §18 migliorano qualità editoriale "
        "e compliance, non dimostrano incremento mandati. Senza fix funnel + calendario owner + "
        "misurazione KPI, il sito resta più una macchina di visibilità informativa che una macchina "
        "di acquisizione incarichi."
    )

    p("—", bold=False)
    p("Documento generato automaticamente dal repository Righetto Immobiliare — audit 29/08/2026.")
    p("Obiettivo commerciale non negoziabile: acquisire nuovi proprietari e nuovi incarichi immobiliari.")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    return OUT_DOCX


def try_pdf(docx_path: Path) -> Path | None:
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(OUT_PDF))
        return OUT_PDF
    except Exception:
        pass
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(OUT_PDF.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
        return OUT_PDF
    except Exception:
        return None


def main() -> int:
    docx = build_docx()
    print(f"OK DOCX: {docx}")
    pdf = try_pdf(docx)
    if pdf and pdf.is_file():
        print(f"OK PDF:  {pdf}")
    else:
        print("PDF: non generato (Word/docx2pdf non disponibile). Apri il DOCX e Salva con nome -> PDF.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
