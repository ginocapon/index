# -*- coding: utf-8 -*-
"""
Genera modello Word — Mandato esclusivo vendita
Gruppo Immobiliare Righetto · Padova 2026

Output: documenti/Modello-Mandato-Esclusivo-Vendita-Righetto-2026.docx
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "documenti" / "Modello-Mandato-Esclusivo-Vendita-Righetto-2026.docx"

F = "[……………………………………………………………………………………]"


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.color.rgb = RGBColor(0x15, 0x24, 0x35)


def add_center(doc: Document, text: str, *, bold: bool = False, size: int = 14) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"


def add_body(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


def add_mixed(doc: Document, parts: list[tuple[str, bool]]) -> None:
    p = doc.add_paragraph()
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)


def add_article(doc: Document, num: str, title: str) -> None:
    doc.add_heading(f"{num}) {title}", level=2)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    style_doc(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    today = date.today().strftime("%d/%m/%Y")

    add_center(doc, "GRUPPO IMMOBILIARE RIGHETTO", bold=True, size=16)
    add_center(doc, "di Capon Gino", size=12)
    add_center(doc, "Via Roma 96 · 35010 Limena (PD)", size=11)
    add_center(doc, "Tel. 049.8843484 · info@righettoimmobiliare.it", size=11)
    doc.add_paragraph()
    add_center(doc, "CONFERIMENTO DI INCARICO DI MEDIAZIONE", bold=True, size=13)
    add_center(doc, "(Vendita — esclusiva)", bold=True, size=12)
    doc.add_paragraph()

    add_body(doc, "La parte sottoscritta:")
    add_body(
        doc,
        f"Il sig./la sig.ra {F}, nato/a a {F} il {F}, C.F. {F}, residente in {F}, "
        f"n. {F} ({F}), identificato/a mediante {F}, n. {F}, rilasciato il {F} con scadenza {F}, "
        f"tel. {F}, e-mail {F},",
    )
    add_body(doc, "di seguito denominato/a «Venditore»,")
    add_body(
        doc,
        "dichiarando di avere e/o rappresentare l'intera proprietà dell'immobile (o, in caso di "
        "comproprietà, di agire con i poteri e le autorizzazioni necessarie per conto di tutti i "
        "comproprietari, come da elenco allegato o da atti di provenienza),",
    )
    add_mixed(
        doc,
        [
            ("CONFERISCE INCARICO DI MEDIAZIONE IN ESCLUSIVA ", True),
            (
                "all'agenzia immobiliare «GRUPPO IMMOBILIARE RIGHETTO di Capon Gino», di seguito "
                "«AI» (Agente Immobiliare), affinché procuri un acquirente per l'immobile sotto "
                "descritto, alle condizioni di seguito indicate.",
                False,
            ),
        ],
    )
    doc.add_paragraph()

    add_article(doc, "1", "Descrizione immobile / dichiarazioni del Venditore")
    add_body(doc, f"Tipo: {F}")
    add_body(doc, f"Ubicazione: {F}")
    add_body(doc, f"Destinazione d'uso: {F}")
    add_body(
        doc,
        f"Descrizione catastale: censito al N.C.E.U. di categoria {F}, comune catastale {F}, "
        f"sezione {F}, foglio {F}, mappale {F}, subalterno {F}, classe {F}, rendita € {F}, "
        f"valore catastale € {F}, consistenza {F}.",
    )
    add_body(doc, f"Composizione: {F}")
    add_body(
        doc,
        f"Proprietà intestata a: {F} (con indicazione delle quote di comproprietà, ove applicabile).",
    )
    add_body(doc, f"Stato dell'immobile: {F}")
    add_body(doc, f"Vincoli / formalità gravanti: {F}")
    add_body(doc, f"Atto di provenienza: {F}")
    add_body(doc, f"ACE/APE, licenza di costruzione, abitabilità, condoni, allineamento catastale: {F}")
    add_body(doc, f"Spese condominiali: {F}")

    add_article(doc, "2", "Prezzo richiesto")
    add_body(doc, f"€ {F} ({F})")
    add_body(doc, "Il saldo prezzo dovrà essere interamente corrisposto al più tardi al momento del rogito notarile.")

    add_article(doc, "3", "Compenso di mediazione")
    add_body(
        doc,
        "Si riconosce all'AI la provvigione pari al 3% (tre percento) più IVA sul prezzo effettivo "
        "di vendita, calcolata sul prezzo di perfezionamento dell'affare.",
    )
    add_mixed(
        doc,
        [
            ("3.1 — Come si divide la provvigione (comproprietà). ", True),
            (
                "In caso di più proprietari, la provvigione totale si ripartisce in base alla "
                "quota di ciascuno (Allegato A o atti). Ogni comproprietario paga solo la propria "
                "parte, per intero: non a rate e non «a metà» della propria quota.",
                False,
            ),
        ],
    )
    add_body(
        doc,
        "Esempio: immobile € 300.000, provvigione 3% = € 9.000 (+ IVA). Tre fratelli con quota "
        "1/3 ciascuno → provvigione totale € 9.000: ciascuno paga € 3.000 (+ IVA sulla propria "
        "quota). Chi firma e conclude la vendita non paga la provvigione degli altri.",
    )
    add_mixed(
        doc,
        [
            ("3.2 — Quando matura. ", True),
            (
                "La provvigione matura al momento in cui l'acquirente viene a conoscenza "
                "dell'accettazione della proposta di acquisto.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("3.3 — Quando è comunque dovuta. ", True),
            (
                "La provvigione resta dovuta se la vendita si perfeziona per effetto dell'attività "
                "dell'AI — anche dopo la scadenza del mandato, se l'acquirente era stato contattato "
                "dall'AI, ovvero se il proponente viene sostituito da altri al rogito.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("3.4 — Acquirente al prezzo del mandato: tutti devono collaborare. ", True),
            (
                "Se l'agenzia ha trovato un acquirente disposto a pagare almeno il prezzo indicato "
                "al punto 2 (proposta accettata, preliminare o altro accordo), tutti i comproprietari "
                "devono cooperare per concludere la vendita. Non basta che uno solo voglia vendere: "
                "serve il consenso di tutti per il rogito.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Chi si rifiuta senza motivo valido ", True),
            (
                "(senza firmare preliminare o rogito pur avendo un acquirente al prezzo di mandato) "
                "resta personalmente obbligato a corrispondere all'AI, a titolo di penale, l'importo "
                "pari al 100% della provvigione totale calcolata sul prezzo di quell'affare (11/11 "
                "della provvigione pattuita). Chi invece ha collaborato e firmato non paga la quota "
                "altrui: paga solo la propria parte di provvigione di cui al comma 3.1, se e quando "
                "la vendita si conclude.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("3.5 — Prezzo proposto più basso del mandato. ", True),
            (
                "Se l'offerta o l'accordo prevedono un prezzo inferiore a quello del punto 2 "
                "(es. mandato € 300.000 e offerta € 280.000), ogni comproprietario è libero di non "
                "firmare e di ridiscutere l'offerta con l'agenzia e con gli altri. In tal caso non "
                "scattano penali e non è dovuta provvigione solo per il fatto del rifiuto.",
                False,
            ),
        ],
    )
    add_body(
        doc,
        "Motivo valido per opporsi (solo a prezzo pieno): vincoli legali oggettivi, forza maggiore "
        "documentata, fatti gravi emersi dopo l'accettazione e non imputabili al venditore. Non è "
        "motivo valido il semplice ripensamento («ci ho ripensato») quando l'offerta rispetta prezzo "
        "e condizioni del mandato.",
    )

    add_article(doc, "4", "Durata dell'incarico")
    add_body(doc, f"Dal {F} al {F}.")
    add_body(
        doc,
        "Per tale periodo l'incarico è irrevocabile. Il presente incarico si rinnoverà tacitamente, "
        "per una sola volta, per uguale periodo, salvo disdetta da parte di uno dei contraenti da "
        "comunicarsi a mezzo raccomandata A/R, fax, e-mail o PEC entro quindici giorni dalla scadenza "
        "dell'incarico.",
    )

    add_article(doc, "5", "Proposta d'acquisto e contratto preliminare")
    add_body(
        doc,
        "L'AI è autorizzata a raccogliere proposte d'acquisto da aspiranti acquirenti, nonché a "
        "trattenere a titolo di deposito infruttifero assegni intestati alla parte venditrice, e a "
        "restituire i medesimi assegni in caso di mancata conclusione dell'affare per cause non "
        "imputabili all'acquirente proponente.",
    )
    add_mixed(
        doc,
        [
            ("Preliminare. ", True),
            (
                "L'AI è autorizzata ad assistere le parti nella stipula del contratto preliminare di "
                "compravendita. Il Venditore si impegna a rispettare gli obblighi assunti verso "
                "l'acquirente con tale atto e a cooperare per il perfezionamento del rogito "
                "definitivo nei termini concordati.",
                False,
            ),
        ],
    )

    add_article(doc, "6", "Esclusiva")
    add_body(
        doc,
        "L'incarico viene conferito in esclusiva, con divieto di conferire a terzi analogo incarico "
        "per tutta la durata dello stesso e impegno a vendere l'immobile soltanto a soggetti messi "
        "in relazione dall'AI.",
    )
    add_body(
        doc,
        "La violazione dell'obbligo di esclusiva, sia nel caso di conferimento di incarico ad altra "
        "agenzia sia per il caso di vendita diretta, comporterà il pagamento da parte del Venditore "
        "della somma prevista a titolo di penale nel successivo art. 7.",
    )
    add_body(
        doc,
        "In conseguenza dell'esclusività e dell'irrevocabilità dell'incarico, l'AI si impegna, con la "
        "propria organizzazione aziendale, alla promozione della vendita dell'immobile, accollandosi "
        "le relative spese di promozione ordinarie e rinunciando al relativo rimborso, salvo diversa "
        "pattuizione scritta.",
    )

    add_article(doc, "7", "Clausola penale")
    add_body(
        doc,
        "A seguito di specifica trattativa individuale, il Venditore che violi gli obblighi sotto "
        "indicati corrisponderà all'AI una penale pari al 100% della provvigione totale pattuita "
        "(art. 3), calcolata sul prezzo dell'affare, nei casi seguenti:",
    )
    add_bullets(
        doc,
        [
            "impedire le visite all'immobile, anche oltre quanto previsto all'art. 8;",
            "far mancare la vendita per informazioni errate o incomplete sull'immobile (gravami, "
            "irregolarità urbanistiche, ipoteche non comunicate);",
            "vendere l'immobile bypassando l'agenzia, in violazione dell'esclusiva;",
            "rifiutare di firmare (preliminare o rogito) pur essendoci un acquirente al prezzo del "
            "mandato, nei casi di cui all'art. 3.4 — con obbligo del solo comproprietario "
            "inadempiente di pagare la penale pari all'intera provvigione; fermo il diritto di non "
            "firmare se il prezzo è inferiore al mandato (art. 3.5);",
            "rifiutare una proposta conforme al mandato e al prezzo richiesto, salvo motivo valido "
            "documentato.",
        ],
    )
    add_body(
        doc,
        "La penale di cui al punto d) non grava sui comproprietari che hanno collaborato e firmato: "
        "ciascuno di questi, in caso di vendita conclusa, paga solo la propria quota di provvigione "
        "(art. 3.1).",
    )
    add_body(
        doc,
        "Viene pattuita una penale pari alla somma pattuita a titolo di provvigione, a carico "
        "dell'AI e a favore del Venditore, nei seguenti casi:",
    )
    add_bullets(
        doc,
        [
            "rinuncia anticipata dell'incarico da parte dell'AI rispetto alla sua naturale scadenza, "
            "senza giustificato motivo;",
            "mancata comunicazione al Venditore di proposte di acquisto conformi al presente "
            "incarico entro un termine ragionevole dalla loro ricezione.",
        ],
    )

    add_article(doc, "8", "Modalità operative")
    add_mixed(
        doc,
        [
            ("Visite e pubblicità. ", True),
            (
                "Si autorizza l'AI a far visionare l'immobile a terzi e a compiere ogni forma di "
                "pubblicità ritenuta idonea all'espletamento dell'incarico, ivi compresa l'esposizione "
                "in loco di cartelli promozionali, annunci su portali e quotidiani, strumenti "
                "informatici, banche dati e collaborazioni con colleghi.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Visite e accesso — giorno settimanale e comproprietario residente. ", True),
            (
                "Dalla sottoscrizione del presente contratto, il Venditore riconosce all'AI la "
                "facoltà di organizzare sopralluoghi con potenziali acquirenti per almeno un giorno "
                "alla settimana, da concordare con l'AI all'avvio dell'incarico, in fasce orarie "
                "comunicate con ragionevole preavviso. Qualora uno o più comproprietari abbiano la "
                "residenza nell'immobile, essi garantiscono la disponibilità in tale giorno, salvo "
                "diverso accordo scritto con il comproprietario interessato. Qualora in tale giorno "
                "il comproprietario residente non possa garantire l'accesso per giusta causa (a "
                "titolo esemplificativo: impossibilità oggettiva, emergenza sanitaria, obblighi "
                "lavorativi non prevedibili), ne darà tempestiva comunicazione all'AI, che "
                "concorderà un nuovo appuntamento con l'acquirente interessato. Il comproprietario "
                "residente che, senza giustificato motivo, renda impossibile o ostacoli in modo "
                "reiterato i sopralluoghi risponde di inadempimento contrattuale, con applicazione "
                "di quanto previsto agli artt. 7, lett. a), e 3, ove compatibile.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Documentazione catastale e urbanistica. ", True),
            (
                "Dalla sottoscrizione dell'accordo con l'acquirente (proposta accettata, preliminare "
                "o atto equivalente), il Venditore si obbliga a mettere a disposizione dell'AI e "
                "dell'acquirente, entro un termine concordato, tutta la documentazione catastale, "
                "urbanistica ed edilizia necessaria alla corretta definizione dell'affare, completa "
                "e in ordine.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Consegna dell'immobile. ", True),
            (
                "Decorso il contratto preliminare, la consegna dell'immobile all'acquirente "
                "avverrà entro e non oltre sei (6) mesi dalla stipula del preliminare stesso, salvo "
                "diverso accordo scritto tra venditore, acquirente e AI.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Inadempimento del comproprietario residente — spese e preliminare. ", True),
            (
                "Il comproprietario che abbia la residenza nell'immobile e, per motivi non previsti "
                "nel presente contratto né giustificati ai sensi dell'art. 3, non consenta la "
                "liberazione dell'immobile o l'accesso per visite e consegna nei termini concordati, "
                "resterà responsabile delle spese, dei danni e degli oneri consequenziali derivanti "
                "dagli accordi con futuri acquirenti e con l'AI. Ove sia stato stipulato un "
                "contratto preliminare e l'inadempimento sia imputabile al Venditore residente, "
                "quest'ultimo — nei limiti di legge e degli obblighi assunti nel preliminare — "
                "dovrà far fronte alle conseguenze previste dal contratto preliminare stesso, ivi "
                "compresa, ove applicabile, la restituzione della caparra confirmatoria in misura "
                "doppia all'acquirente, nonché al compenso di mediazione maturato e dovuto all'AI "
                "ai sensi dell'art. 3.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Disponibilità, regolarità urbanistica e gravami. ", True),
            (
                "Si dichiara di avere la piena disponibilità dell'immobile e che lo stesso è conforme "
                "alla normativa urbanistica ed edilizia vigente, con impianti a norma di legge e munito "
                "di APE ove obbligatorio, con impegno a fornire tempestivamente la documentazione "
                "necessaria.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Comunicazione di ipoteche e formalità pregiudizievoli. ", True),
            (
                "Ciascun comproprietario si obbliga a comunicare tempestivamente all'AI l'esistenza "
                "di ipoteche, pignoramenti, trascrizioni pregiudizievoli o altre formalità gravanti "
                "sulla propria quota di proprietà, al fine di consentire all'AI di informare "
                "l'acquirente e di gestire le pratiche utili alla definizione dell'affare (a titolo "
                "esemplificativo: coordinamento con l'istituto di credito per la liquidazione o il "
                "contemperamento della pendenza). L'omissione dolosa o colposamente negligente di "
                "tali circostanze potrà integrare le fattispecie di cui agli artt. 3 e 7.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Stato dell'immobile alla vendita. ", True),
            (
                "Si garantisce che l'immobile verrà trasferito al rogito libero da oneri e "
                "trascrizioni pregiudizievoli non dichiarate, salvo quelli espressamente accettati "
                "per iscritto dall'acquirente, e comunque regolarizzato sotto il profilo "
                "urbanistico-catastale nei limiti concordati con l'AI e documentati agli atti.",
                False,
            ),
        ],
    )

    add_article(doc, "9", "Doveri del mediatore")
    add_body(
        doc,
        "L'AI, attraverso la propria organizzazione, si impegna ad assistere le parti dal "
        "conferimento dell'incarico fino al rogito notarile di vendita e, considerata l'esclusività "
        "e la durata pattuita, si assume il carico delle spese promozionali ordinarie attive a "
        "promuovere la vendita, svolgendo l'incarico con professionalità e diligenza (art. 1754 c.c.).",
    )

    add_article(doc, "10", "Privacy e formalità obbligatorie")
    add_body(
        doc,
        "Il Venditore dichiara, ai sensi dell'art. 13 del Regolamento UE 2016/679 e del D.lgs. "
        "196/2003, di essere stato informato che i dati personali raccolti saranno trattati ai fini "
        "dell'espletamento del presente incarico. Presto / non presto (barrare) assenso al "
        "trattamento per finalità ulteriori rispetto all'esecuzione contrattuale.",
    )
    add_body(
        doc,
        "Si dichiara di essere edotti delle norme antiriciclaggio (D.lgs. 231/2007 e s.m.i.) e "
        "dell'obbligo di fornire al mediatore la documentazione necessaria per la registrazione "
        "degli atti negoziali nei termini di legge.",
    )

    add_article(doc, "11", "Clausole vessatorie")
    add_body(
        doc,
        "A mente del Codice del Consumo (D.lgs. 206/2005), le parti dichiarano che le clausole del "
        "presente contratto sono state illustrate e oggetto di specifica trattativa individuale, in "
        "particolare quelle relative all'esclusiva, alla divisione della provvigione (art. 3.1), "
        "al rifiuto di firma a prezzo pieno (art. 3.4), all'offerta sotto prezzo (art. 3.5), "
        "al preliminare, alle visite del residente, alla documentazione, alla consegna entro "
        "sei mesi e alle penali di cui all'art. 7, la cui approvazione è espressa in piena "
        "coscienza.",
    )
    add_body(
        doc,
        "Ai sensi degli artt. 1341 e 1342 c.c., le parti approvano specificamente i patti di cui "
        "agli artt. 3, 4, 6, 7, 8, 10 e 11.",
    )

    page_break(doc)

    add_body(doc, f"Limena, lì {F}")
    doc.add_paragraph()
    add_body(doc, "Firma/e del Venditore", bold=True)
    add_body(doc, "_________________________________________________________")
    doc.add_paragraph()
    add_body(doc, "Firma del mediatore per accettazione dell'incarico", bold=True)
    add_body(doc, "Gruppo Immobiliare Righetto di Capon Gino")
    add_body(doc, "_________________________________________________________")
    doc.add_paragraph()
    add_body(doc, f"Limena, lì {F}")
    doc.add_paragraph()
    add_body(
        doc,
        "A mente dell'art. 1341-1342 del Codice civile, le parti specificamente approvano i patti "
        "di cui agli articoli 3, 4, 6, 7, 8, 10 e 11 del presente contratto.",
    )
    doc.add_paragraph()
    add_body(doc, "Firma/e del Venditore", bold=True)
    add_body(doc, "_________________________________________________________")
    doc.add_paragraph()
    add_body(doc, "Firma del mediatore per accettazione dell'incarico", bold=True)
    add_body(doc, "Gruppo Immobiliare Righetto di Capon Gino")
    add_body(doc, "_________________________________________________________")
    doc.add_paragraph()
    add_body(
        doc,
        "Il presente formulario è stato depositato/approvato presso la competente CCIAA di Padova.",
    )
    add_body(doc, "Il presente incarico è stato sottoscritto presso i locali dell'AI.")
    doc.add_paragraph()

    page_break(doc)

    doc.add_heading("Allegato A — Elenco comproprietari e quote", level=1)
    add_body(doc, "Immobile: " + F)
    add_body(doc, "Indirizzo: " + F)
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["N.", "Nome e cognome", "Quota %", "Firma"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in range(1, 5):
        table.rows[row].cells[0].text = str(row)
        for col in range(1, 4):
            table.rows[row].cells[col].text = ""
    doc.add_paragraph()
    add_body(
        doc,
        f"Modello interno Righetto Immobiliare — revisione {today}. Documento di lavoro: verificare "
        "con consulenza legale/notarile prima dell'uso operativo. Deposito CCIAA Padova.",
        bold=False,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    agg = OUT.parent / "Modello-Mandato-Esclusivo-Vendita-Righetto-2026-agg.docx"
    doc.save(str(agg))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"OK: {path}")
