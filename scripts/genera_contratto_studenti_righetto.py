# -*- coding: utf-8 -*-
"""
Genera modello Word — Locazione studenti canone libero 4+4
Gruppo Immobiliare Righetto · Padova 2026

Output: documenti/Modello-Locazione-Studenti-Righetto-2026.docx
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "documenti" / "Modello-Locazione-Studenti-Righetto-2026.docx"


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.color.rgb = RGBColor(0x15, 0x24, 0x35)


def add_center(doc: Document, text: str, *, bold=False, size=14) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"


def add_body(doc: Document, text: str, *, bold=False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.name = "Times New Roman"


def add_article(doc: Document, num: str, title: str, body: str) -> None:
    doc.add_heading(f"Art. {num} — {title}", level=3)
    add_body(doc, body)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    style_doc(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    today = date.today().strftime("%d/%m/%Y")

    # --- COPERTINA ---
    add_center(doc, "GRUPPO IMMOBILIARE RIGHETTO", bold=True, size=18)
    add_center(doc, "Via Roma 96 · 35010 Limena (PD)", size=12)
    add_center(doc, "Tel. 049.8843484 · info@righettoimmobiliare.it", size=12)
    add_center(doc, "righettoimmobiliare.it", size=12)
    doc.add_paragraph()
    add_center(doc, "MODELLO ESCLUSIVO", bold=True, size=16)
    add_center(doc, "Contratto di locazione ad uso abitativo", bold=True, size=14)
    add_center(doc, "Canone libero (4+4) — Singola camera / studenti universitari", bold=True, size=14)
    add_center(doc, "Comune di Padova e provincia", size=12)
    doc.add_paragraph()
    add_center(doc, f"Versione {today}", size=11)
    add_center(doc, "Documento di lavoro — revisione legale consigliata prima dell'uso", size=10)
    page_break(doc)

    # --- INDICE ---
    doc.add_heading("Indice", level=1)
    indice = [
        "1. Nota metodologica e fonti normative verificate",
        "2. Analisi strategica: perché non usare un «contratto 4+4 studenti» tipizzato",
        "3. Proposta professionale Gruppo Immobiliare Righetto",
        "4. Struttura del pacchetto contrattuale (18–20 pagine)",
        "5. Istruzioni per registrazione RLI (Agenzia delle Entrate)",
        "6. CONTRATTO DI LOCAZIONE — testo integrale",
        "7. Allegato A — Planimetria camera locata e spazi comuni",
        "8. Allegato B — Opzione cedolare secca (scheda informativa)",
        "9. Allegato C — Regolamento interno dell'appartamento",
        "10. Allegato D — Ripartizione spese e utenze",
        "11. Allegato E — Inventario arredi (camera e parti comuni)",
        "12. Allegato F — Verbale di consegna",
        "13. Allegato G — Verbale di riconsegna",
        "14. Allegato H — Informativa privacy (art. 13 GDPR)",
        "15. Allegato I — Modulo comunicazione recesso / disdetta",
        "16. Allegato L — Dichiarazioni studente fuori sede (facoltativo)",
        "17. Checklist operativa agenzia",
        "18. Disclaimer e revisione",
    ]
    for row in indice:
        add_body(doc, row)
    page_break(doc)

    # --- SEZIONE 1 ---
    doc.add_heading("1. Nota metodologica e fonti normative verificate", level=1)
    add_body(
        doc,
        "Il presente modello è stato predisposto da Gruppo Immobiliare Righetto sulla base "
        "della normativa vigente e della documentazione ufficiale del Comune di Padova, "
        "senza inventare tipologie contrattuali non previste dalla legge. "
        "Ogni dato normativo riportato è riconducibile a fonte istituzionale.",
    )
    add_body(doc, "Fonti verificate (aggiornamento luglio 2026):", bold=True)
    add_bullets(
        doc,
        [
            "Legge 9 dicembre 1998, n. 431 — Disciplina delle locazioni e del rilascio degli immobili adibiti ad uso abitativo.",
            "Legge 27 luglio 1978, n. 392 — Disciplina delle locazioni abitative (disdetta, cauzione, ecc.).",
            "Comune di Padova — Contratti di locazione (contratti tipo): https://www.comune.padova.it/contratti-di-locazione-contratti-tipo (ultimo aggiornamento pagina: 7/12/2024).",
            "Accordo territoriale per le locazioni nella città di Padova — sottoscritto il 14 febbraio 2023 (protocollo n. 0086231), ai sensi del D.M. 16 gennaio 2017.",
            "D.L. 6 dicembre 2011, n. 201 (c.d. «Salva Italia») — istituzione cedolare secca sugli affitti (art. 3 e ss.).",
            "Regolamento (UE) 2016/679 (GDPR) e D.Lgs. 196/2003 come modificato — trattamento dati personali.",
            "Agenzia delle Entrate — registrazione contratti di locazione (modello RLI telematico).",
        ],
    )
    page_break(doc)

    # --- SEZIONE 2 ---
    doc.add_heading("2. Analisi strategica", level=1)
    doc.add_heading("Perché non esiste un «contratto 4+4 per studenti» tipizzato a Padova", level=2)
    add_body(
        doc,
        "Dopo verifica della normativa comunale, si conferma che il Comune di Padova — "
        "in attuazione dell'Accordo territoriale del 14 febbraio 2023 — prevede tre modelli "
        "standard a canone concordato (con fasce min/max):",
    )
    add_bullets(
        doc,
        [
            "Contratto concertato ad uso abitativo (3+2, estendibile a 4+2, 5+2, 6+2).",
            "Contratto ad uso transitorio (da 1 a 18 mesi, non rinnovabile).",
            "Contratto per studenti universitari dell'Università degli Studi di Padova (durata da 6 a 36 mesi, canone concordato).",
        ],
    )
    add_body(
        doc,
        "Tali schemi (Allegati C, D, E dell'Accordo) si applicano alla locazione a "
        "canone concordato, non al libero mercato.",
    )
    doc.add_heading("La strada corretta per il canone libero", level=2)
    add_body(
        doc,
        "Se si intende locare a canone libero, la forma contrattuale corretta è il "
        "contratto ordinario di locazione ad uso abitativo con durata minima di 4 anni "
        "e rinnovo automatico di ulteriori 4 anni, come previsto dalla legge 431/1998 "
        "per i contratti di libero mercato (cfr. pagina istituzionale Comune di Padova: "
        "«CONTRATTI DI LIBERO MERCATO — durata obbligatoria di 4 anni + 4 di rinnovo»).",
    )
    add_body(
        doc,
        "Il contratto resta pienamente valido e utilizzabile dallo studente per pratiche "
        "universitarie, richiesta alloggio, detrazioni fiscali (ove spettanti) e "
        "iscrizione anagrafica (residenza o domicilio), purché contenga gli elementi "
        "essenziali: identificazione dell'immobile e della camera, durata, canone, "
        "registrazione presso l'Agenzia delle Entrate, consegna dell'Attestato di "
        "Prestazione Energetica ove dovuto.",
    )
    page_break(doc)

    # --- SEZIONE 3 — PROPOSTA (testo potenziato) ---
    doc.add_heading("3. Proposta professionale Gruppo Immobiliare Righetto", level=1)
    add_body(doc, "Sì, e ti consigliamo di procedere in questo modo.", bold=True)
    add_body(
        doc,
        "Riteniamo opportuno predisporre un modello professionale di circa 18–20 pagine, "
        "completamente personalizzato per Gruppo Immobiliare Righetto, che diventi uno "
        "strumento esclusivo per tutte le locazioni universitarie future nel territorio "
        "padovano, più completo dei modelli generici diffusi dalle associazioni di categoria.",
    )
    doc.add_heading("Contenuti del pacchetto", level=2)
    add_bullets(
        doc,
        [
            "Contratto di locazione ad uso abitativo 4+4 (canone libero) con clausole dedicate alla locazione per studenti.",
            "Opzione informativa per l'adesione alla cedolare secca (regime fiscale alternativo per il locatore).",
            "Contratto riferito alla singola camera con uso esclusivo della stanza e uso comune di cucina, bagni, ingresso e altri spazi comuni.",
            "Planimetria allegata con identificazione della camera locata (Allegato A).",
            "Permanenza minima convenuta di 12 mesi, con disciplina del recesso anticipato.",
            "Preavviso di recesso convenuto di 3 mesi (ove compatibile con la disciplina legale applicabile).",
            "Responsabilità solidale se il contratto è unico per più studenti, oppure clausola di autonomia se si stipula un contratto per ogni stanza.",
            "Regolamento interno dell'appartamento (pulizie, raccolta differenziata, ospiti, rumori, manutenzione, spazi comuni).",
            "Modalità di ripartizione delle spese condominiali e utenze.",
            "Divieto di sublocazione e cessione del contratto.",
            "Deposito cauzionale (nei limiti di legge).",
            "Inventario degli arredi.",
            "Verbale di consegna e verbale di riconsegna.",
            "Allegato privacy (art. 13 GDPR).",
            "Modulo per comunicazione del recesso.",
        ],
    )
    doc.add_heading("Clausole operative ad alta efficacia (best practice)", level=2)
    add_bullets(
        doc,
        [
            "Obbligo di comunicare al locatore l'eventuale nuovo coinquilino in caso di sostituzione dello studente.",
            "Divieto di ospitare persone per periodi prolungati senza autorizzazione scritta del locatore.",
            "Obbligo di mantenere pulite le parti comuni e la camera locata.",
            "Possibilità per il locatore di effettuare visita semestrale dell'immobile, previo preavviso di almeno 7 giorni.",
            "Obbligo di riconsegna della camera tinteggiata e pulita (salvo normale usura), pena addebito forfettario.",
            "Disciplina specifica per la restituzione del deposito cauzionale entro 30 giorni dalla riconsegna, con eventuale trattenuta motivata.",
            "Consegna copia dell'APE al conduttore alla sottoscrizione (obbligo di legge).",
            "Registrazione del contratto tramite modello RLI entro 30 giorni (art. 4 D.L. 251/2007).",
        ],
    )
    add_body(
        doc,
        "Il modello è predisposto per essere registrabile tramite RLI dell'Agenzia delle Entrate "
        "e conforme alla normativa vigente, mantenendo la flessibilità del canone libero.",
    )
    page_break(doc)

    # --- SEZIONE 4 ---
    doc.add_heading("4. Struttura del pacchetto contrattuale", level=1)
    add_body(
        doc,
        "Il documento Word che state consultando contiene il testo integrale del contratto "
        "e di tutti gli allegati operativi. In sede di stipula, ogni campo evidenziato con "
        "parentesi quadre […] dovrà essere compilato con i dati dell'operazione specifica.",
    )
    page_break(doc)

    # --- SEZIONE 5 RLI ---
    doc.add_heading("5. Istruzioni per registrazione RLI", level=1)
    add_bullets(
        doc,
        [
            "Accedere al portale «Locazioni» dell'Agenzia delle Entrate (servizi telematici per privati o intermediari abilitati).",
            "Compilare il modello RLI indicando: dati locatore e conduttore, codice fiscale, indirizzo immobile, dati catastali, canone annuo, data inizio e fine (o durata 4 anni), oneri accessori.",
            "Per locazione di singola camera in appartamento condiviso: indicare nell'oggetto la camera e l'indirizzo completo dell'unità; allegare planimetria come da prassi locale.",
            "Versare l'imposta di registro (2% del canone annuo, salvo agevolazioni) e le bolli ove dovuti.",
            "Conservare ricevuta di registrazione e codice RLI da comunicare alle parti.",
            "Termine: entro 30 giorni dalla stipula (art. 4, comma 1, D.L. 25 novembre 2007, n. 251).",
        ],
    )
    page_break(doc)

    # --- CONTRATTO ---
    doc.add_heading("6. CONTRATTO DI LOCAZIONE AD USO ABITATIVO", level=1)
    add_center(doc, "(Canone libero — singola camera — studenti universitari)", size=12)
    doc.add_paragraph()

    add_article(
        doc,
        "1",
        "Parti",
        "Tra il Sig./la Sig.ra [NOME COGNOME LOCATORE], nato/a a […] il […], C.F. […], "
        "residente in […], di seguito «Locatore», e il/la Sig./Sig.ra [NOME COGNOME CONDUTTORE], "
        "nato/a a […] il […], C.F. […], studente/a iscritto/a al […] anno del corso di "
        "[…] presso l'Università degli Studi di […], matricola n. […], di seguito «Conduttore», "
        "si stipula il presente contratto di locazione ad uso abitativo.",
    )
    add_article(
        doc,
        "2",
        "Premesse",
        "a) Il Locatore è proprietario / legittimo detentore dell'unità immobiliare sita in "
        "Padova (o comune di […]), Via/Piazza […] n. […], piano […], interno […], "
        "identificata al N.C.E.U. al foglio […], mappale/i […], sub. […].\n"
        "b) L'immobile è dotato di Attestato di Prestazione Energetica (APE) classe […], "
        "protocollo n. […], consegnato in copia al Conduttore.\n"
        "c) Le parti convengono la locazione a canone libero, ai sensi della legge 431/1998, "
        "con durata minima quadriennale e successivo rinnovo quadriennale, integrando il "
        "presente atto con clausole specifiche per la locazione a studenti universitari.\n"
        "d) Il Conduttore dichiara di destinare la camera locata a propria abitazione, "
        "anche ai fini dell'iscrizione anagrafica (residenza o domicilio), ove ne ricorrano i presupposti.",
    )
    add_article(
        doc,
        "3",
        "Oggetto",
        "Il Locatore concede in locazione al Conduttore, che accetta, la camera contrassegnata "
        "con il n. […] (Allegato A), nonché l'uso comune, in proporzione agli altri conduttori "
        "presenti, di cucina, bagno/i, ingresso, corridoio, balcone/terrazzo […] e degli impianti "
        "comuni dell'appartamento. È escluso dall'uso del Conduttore ogni altro locale non "
        "espressamente indicato. La camera ha una superficie di circa […] mq ed è arredata come "
        "da inventario (Allegato E).",
    )
    add_article(
        doc,
        "4",
        "Durata",
        "Il contratto ha durata di anni 4 (quattro) a decorrere dal [DATA INIZIO] e si rinnova "
        "automaticamente per ulteriori 4 (quattro) anni, salvo disdetta del Locatore comunicata "
        "con i termini di legge (art. 2, legge 431/1998, e art. 3, legge 392/1978). "
        "Le parti convengono, per la specificità della locazione a studenti, una permanenza "
        "minima del Conduttore di mesi 12 (dodici) dalla data di inizio; il recesso anticipato "
        "prima di tale termine comporta il pagamento di una penale pari a […] mensilità del "
        "canone, salvo concordi diversi e fatti salvi i casi di forza maggiore documentati.",
    )
    add_article(
        doc,
        "5",
        "Canone e pagamento",
        "Il canone mensile, libero, è convenuto in euro […] ([…]/00), da corrispondersi entro "
        "il giorno […] di ogni mese mediante bonifico bancario sul c/c IBAN […] intestato a "
        "[…]. La prima rata sarà corrisposta alla sottoscrizione. Il canone si intende "
        "comprensivo / non comprensivo (barrare) delle spese condominiali ordinarie secondo "
        "Allegato D. Eventuale adeguamento ISTAT: [sì/no — specificare].",
    )
    add_article(
        doc,
        "6",
        "Deposito cauzionale",
        "A garanzia delle obbligazioni del Conduttore, questi versa al Locatore la somma di "
        "euro […] a titolo di deposito cauzionale, pari a n. […] mensilità del canone, "
        "nei limiti massimi di legge (art. 11, legge 392/1978). Il deposito sarà restituito "
        "entro 30 giorni dalla riconsegna delle chiavi, previa verifica dello stato dell'immobile "
        "e del regolare pagamento di canoni, utenze e spese, con eventuale trattenuta motivata "
        "per danni o mancata pulizia/tinteggiatura come da Allegato G.",
    )
    add_article(
        doc,
        "7",
        "Spese e utenze",
        "Le spese condominiali ordinarie, l'acqua, il gas, l'energia elettrica, internet e "
        "tassa rifiuti (TARI) sono ripartite come da Allegato D. Il Conduttore si impegna a "
        "reimburseare la propria quota entro […] giorni dalla richiesta del Locatore o "
        "dell'amministratore, presentando ove possibile i propri consumi.",
    )
    add_article(
        doc,
        "8",
        "Obblighi del Conduttore",
        "Il Conduttore si obbliga a: (a) usare la camera e le parti comuni con diligenza; "
        "(b) osservare il regolamento interno (Allegato C); (c) non sublocare, non cedere il "
        "contratto e non ospitare terzi per più di […] notti consecutive senza autorizzazione "
        "scritta del Locatore; (d) comunicare tempestivamente eventuali guasti; (e) consentire "
        "al Locatore di effettuare un sopralluogo semestrale, previo preavviso di almeno 7 giorni; "
        "(f) in caso di sostituzione del coinquilino/studente, comunicare al Locatore i dati "
        "del nuovo occupante prima dell'ingresso; (g) riconsegnare la camera pulita e tinteggiata "
        "a fine locazione, salvo normale deperimento d'uso.",
    )
    add_article(
        doc,
        "9",
        "Obblighi del Locatore",
        "Il Locatore consegna l'immobile in buono stato di manutenzione e garantisce il "
        "godimento pacifico per la durata del contratto. È responsabile delle riparazioni "
        "straordinarie e degli adempimenti non imputabili al Conduttore, salvo dolo o colpa "
        "di quest'ultimo.",
    )
    add_article(
        doc,
        "10",
        "Manutenzione",
        "Le riparazioni di ordinaria manutenzione a carico del Conduttore sono quelle di "
        "cui all'art. 1609 c.c. e alla prassi locatizia (sostituzione lampadine, piccole "
        "riparazioni d'uso). Quelle straordinarie restano a carico del Locatore.",
    )
    add_article(
        doc,
        "11",
        "Recesso e disdetta",
        "Per il recesso del Conduttore in corso di locazione, le parti convengono un preavviso "
        "minimo di mesi 3 (tre) da comunicare a mezzo raccomandata A/R o PEC a […], fatto salvo "
        "il rispetto dei termini minimi di legge per la disdetta alla scadenza del quadriennio "
        "(art. 3, legge 392/1978: almeno 6 mesi prima della scadenza del periodo in corso). "
        "Il modulo di comunicazione è nell'Allegato I.",
    )
    add_article(
        doc,
        "12",
        "Responsabilità solidale (opzione — barrare se non applicabile)",
        "Qualora più studenti abbiano sottoscritto un unico contratto, ciascuno risponde in "
        "solido verso il Locatore per canoni, spese e obbligazioni. In alternativa, con "
        "contratti separati per ciascuna camera, la responsabilità del presente Conduttore "
        "resta limitata alla propria camera e alla quota parte delle spese comuni.",
    )
    add_article(
        doc,
        "13",
        "Registrazione e imposte",
        "Le spese di registrazione del presente contratto, bolli e imposta di registro, "
        "sono a carico di [Locatore / Conduttore / metà ciascuno]. Le parti si impegnano a "
        "provvedere alla registrazione tramite modello RLI entro 30 giorni dalla stipula.",
    )
    add_article(
        doc,
        "14",
        "Cedolare secca (informativa)",
        "Il Locatore dichiara di [aderire / non aderire] al regime della cedolare secca per "
        "il presente immobile, ai sensi dell'art. 3 del D.L. 6 dicembre 2011, n. 201. "
        "Scheda informativa in Allegato B.",
    )
    add_article(
        doc,
        "15",
        "Privacy",
        "I dati personali delle parti sono trattati secondo l'Allegato H (informativa art. 13 GDPR).",
    )
    add_article(
        doc,
        "16",
        "Foro competente",
        "Per ogni controversia è competente il Foro di Padova, salvo il foro inderogabile "
        "del consumatore ove applicabile.",
    )
    add_article(
        doc,
        "17",
        "Allegati",
        "Formano parte integrante del contratto: Allegati A, C, D, E, F, H, I; verbale di "
        "riconsegna (G) alla cessazione.",
    )
    add_body(doc, "Letto, approvato e sottoscritto.")
    add_body(doc, "Luogo e data: Padova, lì […]")
    add_body(doc, "Il Locatore _________________________")
    add_body(doc, "Il Conduttore _________________________")
    add_body(doc, "(Per esteso, se richiesto dalla registrazione)")
    page_break(doc)

    # --- ALLEGATI ---
    doc.add_heading("7. Allegato A — Planimetria", level=1)
    add_body(doc, "[Inserire planimetria catastale o rilievo interno con evidenziata in colore la camera n. …]")
    add_body(doc, "Legenda: zona verde = camera locata; zona gialla = spazi comuni.")
    page_break(doc)

    doc.add_heading("8. Allegato B — Cedolare secca", level=1)
    add_body(
        doc,
        "La cedolare secca è un regime fiscale opzionale per i redditi da locazione di "
        "immobili ad uso abitativo situati in Italia. Aliquota ordinaria 21%; aliquota 10% "
        "per contratti a canone concordato che rispettano l'Accordo territoriale. "
        "Per il canone libero si applica generalmente l'aliquota del 21%. "
        "L'opzione ha durata triennale e va esercitata nel modello 730/Unico o tramite "
        "intermediario abilitato. Verificare sempre la normativa vigente e la propria situazione fiscale.",
    )
    page_break(doc)

    doc.add_heading("9. Allegato C — Regolamento interno", level=1)
    regolamento = [
        "Orari di quiete: 22:00–08:00 nei giorni feriali; 23:30–09:00 nei week-end.",
        "Pulizia: rotazione settimanale delle parti comuni secondo tabella […]; ogni conduttore mantiene in ordine la propria camera.",
        "Raccolta differenziata: conforme al calendario comunale; smaltimento corretto dei rifiuti.",
        "Ospiti: massimo […] notti/mese; ospiti per periodi superiori solo con autorizzazione scritta.",
        "Rumori e feste: vietate attività che arrechino disturbo agli altri inquilini e al vicinato.",
        "Animali: [ammessi / non ammessi — specificare].",
        "Fumo: vietato all'interno dell'appartamento.",
        "Chiavi: divieto di duplicazione senza consenso; consegna chiavi come da verbale F.",
        "Guasti: comunicazione immediata al Locatore o all'agenzia Gruppo Immobiliare Righetto (049.8843484).",
    ]
    for i, r in enumerate(regolamento, 1):
        add_body(doc, f"{i}. {r}")
    page_break(doc)

    doc.add_heading("10. Allegato D — Ripartizione spese", level=1)
    add_body(doc, "Spese condominiali ordinarie: ripartizione […]% per conduttore / n. occupanti.")
    add_body(doc, "Utenze (luce, gas, acqua): [individuali con sottocontatori / ripartizione a quote uguali / altro …].")
    add_body(doc, "Internet/Wi-Fi: euro […]/mese a carico di […].")
    add_body(doc, "TARI: a carico del Conduttore se contratto > 6 mesi (verificare normativa comunale).")
    page_break(doc)

    doc.add_heading("11. Allegato E — Inventario arredi", level=1)
    add_body(doc, "Camera: letto […], armadio […], scrivania […], sedia […], lampada […], altro […].")
    add_body(doc, "Parti comuni: tavolo […], sedie […], divano […], elettrodomestici cucina […].")
    add_body(doc, "Stato: buono / discreto / da segnalare — note: […]")
    page_break(doc)

    doc.add_heading("12. Allegato F — Verbale di consegna", level=1)
    add_body(doc, "In data […] il Locatore consegna e il Conduttore riceve la camera e le chiavi.")
    add_body(doc, "Lettura contatori: luce […], gas […], acqua […].")
    add_body(doc, "Stato immobile: conforme inventario / difformità: […]")
    add_body(doc, "Firme: Locatore _________ Conduttore _________")
    page_break(doc)

    doc.add_heading("13. Allegato G — Verbale di riconsegna", level=1)
    add_body(doc, "In data […] il Conduttore riconsegna la camera.")
    add_body(doc, "Verifica tinteggiatura, pulizia, danni: […]")
    add_body(doc, "Restituzione deposito cauzionale: euro […] entro 30 giorni.")
    add_body(doc, "Firme: Locatore _________ Conduttore _________")
    page_break(doc)

    doc.add_heading("14. Allegato H — Privacy (art. 13 GDPR)", level=1)
    add_body(
        doc,
        "Titolare del trattamento: [Nome Locatore / Gruppo Immobiliare Righetto quale incaricato]. "
        "Finalità: gestione del rapporto locativo, adempimenti fiscali e registrazione contratto, "
        "contatti per manutenzione. Base giuridica: esecuzione contratto (art. 6.1.b GDPR) e "
        "obblighi di legge (art. 6.1.c). Conservazione: 10 anni dalla cessazione del rapporto "
        "per finalità fiscali. Diritti dell'interessato: accesso, rettifica, cancellazione, "
        "limitazione, reclamo al Garante Privacy. Contatto: info@righettoimmobiliare.it.",
    )
    page_break(doc)

    doc.add_heading("15. Allegato I — Modulo recesso / disdetta", level=1)
    add_body(doc, "Spett.le [Locatore], presente la presente per comunicare [recesso anticipato / disdetta alla scadenza] del contratto di locazione dell'immobile in [indirizzo], con decorrenza dal [data], rispettando il preavviso di mesi […].")
    add_body(doc, "Data, firma del Conduttore, estremi raccomandata A/R o PEC.")
    page_break(doc)

    doc.add_heading("16. Allegato L — Dichiarazioni studente (facoltativo)", level=1)
    add_body(
        doc,
        "Il Conduttore dichiara di essere regolarmente iscritto all'Università […], "
        "di avere residenza anagrafica nel comune di […] (diverso da Padova ove applicabile "
        "per pratiche ESU), e di utilizzare la camera come abitazione principale durante "
        "il periodo accademico.",
    )
    page_break(doc)

    doc.add_heading("17. Checklist operativa agenzia", level=1)
    checklist = [
        "Verificare titolo abitativo e conformità urbanistica dell'immobile.",
        "Verificare APE valido e consegnare copia al conduttore.",
        "Compilare tutti i campi […] e allegati A, C, D, E, F.",
        "Far firmare contratto, regolamento e verbale di consegna.",
        "Incassare cauzione e prima mensilità; documentare pagamenti.",
        "Registrare contratto RLI entro 30 giorni.",
        "Comunicare al conduttore modalità per residenza/domicilio in Comune.",
        "Archiviare copia digitale in fascicolo immobile.",
        "Programmare visita semestrale con preavviso.",
    ]
    for c in checklist:
        add_body(doc, f"☐ {c}")
    page_break(doc)

    doc.add_heading("18. Disclaimer e revisione", level=1)
    add_body(
        doc,
        "Il presente modello è uno strumento di lavoro professionale predisposto da "
        "Gruppo Immobiliare Righetto. Non sostituisce il parere di un legale o di un "
        "consulente fiscale. Prima dell'utilizzo in operazioni reali si raccomanda la "
        "revisione da parte di un avvocato civilista e la verifica della conformità "
        "alle circolari aggiornate dell'Agenzia delle Entrate. "
        "Compenso di mediazione: da concordare in sede con il mandato — non sono riportate "
        "percentuali o listini in questo documento.",
    )
    add_body(doc, f"Documento generato il {today} — Gruppo Immobiliare Righetto.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"OK: {path}")
