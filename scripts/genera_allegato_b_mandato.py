# -*- coding: utf-8 -*-
"""
Genera Allegato B — Scheda informativa comproprietari (Word)
Output: documenti/Allegato-B-Scheda-Informativa-Mandato-Esclusivo-Comproprieta.docx
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "documenti" / "Allegato-B-Scheda-Informativa-Mandato-Esclusivo-Comproprieta.docx"


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)


def add_center(doc: Document, text: str, *, bold: bool = False, size: int = 14) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"


def add_body(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)


def add_heading(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0x2C, 0x4A, 0x6E)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)


def add_box(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)
    r.italic = True


def build() -> Path:
    doc = Document()
    style_doc(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    today = date.today().strftime("%d/%m/%Y")

    add_center(doc, "GRUPPO IMMOBILIARE RIGHETTO", bold=True, size=13)
    add_center(doc, "Allegato B — Scheda informativa per i comproprietari", bold=True, size=12)
    add_center(doc, "Mandato esclusivo di vendita", size=11)
    add_center(
        doc,
        f"Da leggere e firmare prima o insieme al contratto di incarico · revisione {today}",
        size=10,
    )
    doc.add_paragraph()

    add_body(
        doc,
        "Questa pagina riassume, in linguaggio chiaro, i punti principali del mandato. "
        "Il testo completo e vincolante resta quello del contratto di mediazione sottoscritto in agenzia.",
    )
    doc.add_paragraph()

    add_heading(doc, "1. Provvigione — come si divide")
    add_bullets(
        doc,
        [
            "La provvigione dell'agenzia (es. 3% + IVA sul prezzo di vendita) si calcola sul prezzo "
            "effettivo concordato con l'acquirente.",
            "In caso di comproprietà, ogni proprietario paga la propria parte in base alla quota di "
            "proprietà (es. due fratelli al 50% → ciascuno metà della provvigione).",
            "Ognuno è tenuto a corrispondere per intero (11/11) la propria quota: non si può pagare "
            "solo una frazione della propria parte.",
        ],
    )

    add_heading(doc, "2. Se c'è un acquirente al prezzo del mandato")
    add_box(
        doc,
        "In sintesi: se l'agenzia trova un acquirente disposto a pagare almeno il prezzo indicato "
        "nel mandato, tutti i comproprietari devono collaborare per concludere la vendita "
        "(proposta, preliminare, rogito).",
    )
    add_bullets(
        doc,
        [
            "Non si può opporsi senza un motivo serio e documentato (es. vincoli legali, forza "
            "maggiore — non il semplice «ci ripenso»).",
            "Chi si rifiuta di firmare in questa situazione risponde personalmente: deve pagare la "
            "provvigione maturata e la penale prevista dal contratto (pari al 100% / 11/11 della "
            "provvigione pattuita).",
        ],
    )

    add_heading(doc, "3. Se il prezzo proposto è più basso del mandato")
    add_bullets(
        doc,
        [
            "Ognuno resta libero di non accettare e di ridiscutere l'offerta con l'agenzia e con "
            "gli altri proprietari.",
            "In questo caso non scattano penali né obbligo di provvigione per il solo rifiuto.",
        ],
    )

    add_heading(doc, "4. Visite — chi abita in casa")
    add_bullets(
        doc,
        [
            "L'agenzia organizzerà le visite con gli acquirenti per almeno un giorno a settimana, "
            "da concordare all'inizio (giorno e fasce orarie).",
            "Se un comproprietario risiede nell'immobile, deve garantire l'accesso in quel giorno, "
            "salvo diverso accordo scritto.",
            "Se in quel giorno non è possibile per giusta causa (malattia, emergenza, impegno "
            "imprevisto), va comunicato subito all'agenzia per fissare un altro appuntamento.",
            "Bloccare le visite senza motivo valido può comportare penali e inadempimento contrattuale.",
        ],
    )

    add_heading(doc, "5. Documenti, preliminare e consegna")
    add_bullets(
        doc,
        [
            "Dall'accordo con l'acquirente (proposta accettata o preliminare), tutti devono mettere "
            "a disposizione la documentazione catastale e urbanistica completa (planimetrie, APE, "
            "eventuali regolarizzazioni).",
            "L'agenzia può assistere nella stipula del contratto preliminare; chi firma si impegna "
            "a rispettare quanto concordato.",
            "Dopo il preliminare, la consegna dell'immobile all'acquirente deve avvenire entro "
            "6 mesi, salvo diverso accordo scritto tra tutte le parti.",
        ],
    )

    add_heading(doc, "6. Chi vive in casa e non libera l'immobile")
    add_box(
        doc,
        "Se il comproprietario residente, senza motivi previsti dal contratto, non consente "
        "visite, liberazione o consegna nei termini concordati, risponde delle spese e dei danni "
        "che ne derivano verso acquirenti e agenzia.",
    )
    add_bullets(
        doc,
        [
            "Se è già stato firmato un preliminare e l'inadempimento è imputabile a chi abita in "
            "casa, potranno applicarsi anche le conseguenze previste dal preliminare — tra cui, "
            "ove previsto dalla legge e dal contratto con l'acquirente, la restituzione del doppio "
            "della caparra — oltre alla provvigione dovuta all'agenzia.",
        ],
    )

    add_heading(doc, "7. Cosa conviene fare prima di firmare")
    add_bullets(
        doc,
        [
            "Verificare che l'elenco comproprietari e quote (Allegato A) sia corretto.",
            "Concordare subito il giorno settimanale per le visite se qualcuno abita nell'immobile.",
            "Comunicare all'agenzia ipoteche, pignoramenti o altri vincoli sulla propria quota, se presenti.",
            "Chiarire dubbi in agenzia: 049.8843484 · info@righettoimmobiliare.it",
        ],
    )

    doc.add_paragraph()
    add_body(
        doc,
        "Dichiarazione di presa visione — Dichiaro di aver ricevuto e letto la presente scheda "
        "informativa e di aver avuto modo di porre domande in agenzia prima della sottoscrizione "
        "del mandato esclusivo.",
        bold=True,
    )
    doc.add_paragraph()

    for n in range(1, 5):
        add_body(doc, f"Comproprietario {n}", bold=True)
        add_body(doc, "Nome: _________________________________  Data: ___/___/______")
        add_body(doc, "Firma: _________________________________")
        doc.add_paragraph()

    add_body(
        doc,
        f"Documento interno Righetto Immobiliare — revisione {today}. Allegato B informativo, "
        "non sostitutivo del contratto di mediazione. Verificare con consulenza legale/notarile "
        "prima dell'uso operativo.",
        size=9,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"OK: {path}")
